import os
from reportlab.pdfgen import canvas
from docx import Document

FIXTURE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "tests", "fixtures")
os.makedirs(FIXTURE_DIR, exist_ok=True)

def create_pdf(path, lines):
    c = canvas.Canvas(path)
    y = 750
    for line in lines:
        c.drawString(100, y, line)
        y -= 25
    c.showPage()
    c.save()

def generate_fixtures():
    # 1. Native-text PDF: HCPCS 97110, M17.11, State CO, referral, conservative treatment B 1 month. (Missing joint diagnosis)
    create_pdf(
        os.path.join(FIXTURE_DIR, "native_text.pdf"),
        [
            "Clinical Referral Note",
            "Patient: John Doe",
            "DOB: 1970-05-15",
            "Requested Procedure Code: 97110",
            "Diagnosis: M17.11 (Osteoarthritis of right knee joint)",
            "State of Service: CO",
            "Referral from Dr. Smith: referred for therapy services by a physician",
            "Prior treatment: conservative physical therapy treatment B failed after 1 months.",
        ]
    )

    # 2. DOCX: Complete APPROVE evidence
    doc = Document()
    doc.add_heading("Clinical Treatment Summary", 0)
    doc.add_paragraph("Patient: DOCX Approve Patient")
    doc.add_paragraph("DOB: 1954-03-20")
    doc.add_paragraph("Requested CPT Code: 97110")
    doc.add_paragraph("State: CO")
    doc.add_paragraph("Primary Diagnosis Code: M17.11 (Osteoarthritis of right knee joint)")
    doc.add_paragraph("Joint Impairment: Documented joint musculoskeletal impairment of knee.")
    doc.add_paragraph("Referral: referred for therapy services by a physician.")
    doc.add_paragraph("Treatment History: Conservative treatment B failed after 6 months.")
    doc.save(os.path.join(FIXTURE_DIR, "complete_approve.docx"))

    # 3. TXT: NOT_MET evidence (e.g. age = 52 while NCD requires >= 65)
    with open(os.path.join(FIXTURE_DIR, "not_met.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join([
            "CLINICAL ASSESSMENT",
            "Patient: TXT NotMet Patient",
            "DOB: 1974-08-10",
            "Age: 52",
            "Requested HCPCS: 97110",
            "State: CO",
            "Diagnosis: M17.11",
            "Referral: referred for therapy services by a physician",
            "Prior treatment B failed after 1 months."
        ]))

    # 4. Document missing HCPCS / State
    create_pdf(
        os.path.join(FIXTURE_DIR, "missing_params.pdf"),
        [
            "Clinical Assessment Note",
            "Patient: Missing Params Patient",
            "Diagnosis: M17.11 (Osteoarthritis of knee)",
            "Referral: referred for therapy services by a physician",
            "HCPCS: absent",
            "State: absent",
        ]
    )

    # 5. Outcome leakage document
    create_pdf(
        os.path.join(FIXTURE_DIR, "leakage_outcome.pdf"),
        [
            "Payer Case Review",
            "Patient: Leakage Protected Patient",
            "Requested HCPCS: 97110",
            "State: CO",
            "Diagnosis Code: M17.11",
            "Prior authorization approved",
            "Meets policy criteria",
            "Coverage denied",
            "Medical necessity met"
        ]
    )

    # 6. Conflicting evidence document
    create_pdf(
        os.path.join(FIXTURE_DIR, "conflict_doc.pdf"),
        [
            "Clinical Summary Note",
            "Patient: Conflicting Patient",
            "DOB: 1955-03-20", # Conflicts with structured 1954-03-20
            "Requested CPT: 97110",
            "State: CO",
            "Diagnosis Code: M17.11",
            "Prior treatment: Total knee replacement surgery done in 2025." # Conflicts with "no surgery"
        ]
    )

    # 7. Scanned PDF (empty native text)
    create_pdf(os.path.join(FIXTURE_DIR, "scanned_only.pdf"), [])

    print("Test document fixtures generated successfully.")

if __name__ == "__main__":
    generate_fixtures()
